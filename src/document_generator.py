"""Safe Word template placeholder replacement and DOCX generation."""

from __future__ import annotations

import logging
import os
import re
import shutil
import time
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.models import PlaceholderContext
from src.word_process import quit_all_word_applications, remove_file_with_retry

logger = logging.getLogger(__name__)

UNREPLACED_PATTERN = re.compile(r"\{\{[A-Z_]+\}\}")


class DocumentGeneratorError(Exception):
    """Raised when document generation fails."""


class UnreplacedPlaceholdersError(DocumentGeneratorError):
    """Raised when placeholders remain after replacement."""


def generate_docx(
    template_path: Path,
    output_path: Path,
    context: PlaceholderContext,
) -> Path:
    """Copy template, replace placeholders, save DOCX."""
    if not template_path.exists():
        raise DocumentGeneratorError(f"Word template not found: {template_path}")

    output_path = output_path.resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = output_path.with_suffix(output_path.suffix + ".tmp")

    last_os_error: OSError | None = None

    for attempt in range(1, 3):
        try:
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)

            shutil.copy2(template_path, temp_path)
            doc = Document(str(temp_path))
            replacements = context.as_dict()
            warnings = replace_placeholders_in_document(doc, replacements)
            for warning in warnings:
                logger.warning(warning)

            doc.save(str(temp_path))
            remaining = find_unreplaced_in_document(doc)
            if remaining:
                raise UnreplacedPlaceholdersError(
                    f"Unreplaced placeholders in document: {', '.join(sorted(remaining))}"
                )

            _publish_docx(temp_path, output_path)
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)
            return output_path
        except UnreplacedPlaceholdersError:
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)
            raise
        except OSError as exc:
            last_os_error = exc
            logger.warning(
                "DOCX generation attempt %d failed for %s: %s",
                attempt,
                output_path,
                exc,
            )
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)
            if attempt < 2:
                quit_all_word_applications(force_kill=True)
                time.sleep(1)
            else:
                raise DocumentGeneratorError(
                    f"не удалось создать DOCX (файл занят другим процессом): {output_path}"
                ) from exc
        except DocumentGeneratorError:
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)
            raise
        except Exception as exc:
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)
            raise DocumentGeneratorError(f"ошибка создания DOCX: {exc}") from exc

    if last_os_error is not None:
        raise DocumentGeneratorError(f"ошибка создания DOCX: {last_os_error}") from last_os_error
    raise DocumentGeneratorError(f"не удалось создать DOCX: {output_path}")


def _publish_docx(temp_path: Path, output_path: Path) -> None:
    """Move temp DOCX into final location, replacing a locked file if needed."""
    if output_path.exists():
        remove_file_with_retry(output_path, max_attempts=3, quit_word_on_failure=True)
    os.replace(temp_path, output_path)


def replace_placeholders_in_document(
    doc: DocumentObject,
    replacements: dict[str, str],
) -> list[str]:
    """Replace placeholders in body, tables, headers, and footers."""
    warnings: list[str] = []
    found_keys: set[str] = set()

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements, found_keys)

    for table in doc.tables:
        _replace_in_table(table, replacements, found_keys)

    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is None:
                continue
            for paragraph in header_footer.paragraphs:
                _replace_in_paragraph(paragraph, replacements, found_keys)
            for table in header_footer.tables:
                _replace_in_table(table, replacements, found_keys)

    for key in replacements:
        if key not in found_keys:
            warnings.append(f"Placeholder not found in template: {key}")

    return warnings


def _replace_in_table(
    table: Table,
    replacements: dict[str, str],
    found_keys: set[str],
) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements, found_keys)
            for nested in cell.tables:
                _replace_in_table(nested, replacements, found_keys)


def _replace_in_paragraph(
    paragraph: Paragraph,
    replacements: dict[str, str],
    found_keys: set[str],
) -> None:
    if not paragraph.runs:
        text = paragraph.text
        if not text:
            return
        new_text = _apply_replacements(text, replacements, found_keys)
        if new_text != text:
            paragraph.add_run(new_text)
        return

    # Pass 1: replace inside each run — keeps font/size (e.g. Arial 14 greeting).
    for run in paragraph.runs:
        if not run.text:
            continue
        new_text = _apply_replacements(run.text, replacements, found_keys)
        if new_text != run.text:
            run.text = new_text

    # Pass 2: placeholders split across runs (rare).
    full_text = "".join(run.text for run in paragraph.runs)
    if UNREPLACED_PATTERN.search(full_text):
        _replace_cross_run_in_paragraph(paragraph, replacements, found_keys)


def _replace_cross_run_in_paragraph(
    paragraph: Paragraph,
    replacements: dict[str, str],
    found_keys: set[str],
) -> None:
    """Replace placeholders split across runs; formatting of the first involved run is kept."""
    runs = paragraph.runs
    max_iterations = 50

    for _ in range(max_iterations):
        full_text = "".join(run.text for run in runs)
        match = UNREPLACED_PATTERN.search(full_text)
        if not match:
            break

        placeholder = match.group()
        if placeholder not in replacements:
            break

        start, end = match.start(), match.end()
        value = replacements[placeholder]
        found_keys.add(placeholder)

        pos = 0
        start_run = end_run = 0
        start_off = end_off = 0
        for idx, run in enumerate(runs):
            run_len = len(run.text)
            if pos <= start < pos + run_len:
                start_run = idx
                start_off = start - pos
            if pos < end <= pos + run_len:
                end_run = idx
                end_off = end - pos
                break
            pos += run_len

        before = runs[start_run].text[:start_off]
        after = runs[end_run].text[end_off:]
        runs[start_run].text = before + value + after
        for idx in range(start_run + 1, end_run + 1):
            runs[idx].text = ""


def _apply_replacements(
    text: str,
    replacements: dict[str, str],
    found_keys: set[str],
) -> str:
    result = text
    for placeholder, value in replacements.items():
        if placeholder in result:
            found_keys.add(placeholder)
            result = result.replace(placeholder, value)
    return result


def find_unreplaced_in_document(doc: DocumentObject) -> set[str]:
    """Scan an in-memory document for remaining {{...}} placeholders."""
    texts: list[str] = []

    for paragraph in doc.paragraphs:
        texts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    texts.append(paragraph.text)

    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is None:
                continue
            for paragraph in header_footer.paragraphs:
                texts.append(paragraph.text)
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            texts.append(paragraph.text)

    remaining: set[str] = set()
    for text in texts:
        remaining.update(UNREPLACED_PATTERN.findall(text))
    return remaining


def find_unreplaced_placeholders(docx_path: Path) -> set[str]:
    """Scan a DOCX file on disk for remaining {{...}} placeholders."""
    return find_unreplaced_in_document(Document(str(docx_path)))
