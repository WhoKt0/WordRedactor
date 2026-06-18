"""Create long/short greeting template copies from Template_word.docx."""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.template_selector import TEMPLATE_DEFAULT, TEMPLATE_LONG, TEMPLATE_SHORT

_GREETING_MARKER = "{{GREETING_WORD}}"


def _find_greeting_paragraph(doc: Document):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if _GREETING_MARKER in paragraph.text:
                        return paragraph
    return None


def _greeting_run_index(paragraph) -> int | None:
    for index, run in enumerate(paragraph.runs):
        if _GREETING_MARKER in run.text:
            return index
    return None


def _adjust_greeting_indent(doc: Document, *, delta_spaces: int) -> bool:
    """
    Shift leading spaces before {{GREETING_WORD}} without touching run formatting.

    Only whitespace-only runs before the greeting run are modified so Arial 14
    on the greeting placeholders is preserved.
    """
    paragraph = _find_greeting_paragraph(doc)
    if paragraph is None or not paragraph.runs:
        return False

    greeting_run_idx = _greeting_run_index(paragraph)
    if greeting_run_idx is None:
        return False

    if delta_spaces == 0:
        return True

    if delta_spaces > 0:
        if greeting_run_idx > 0:
            paragraph.runs[0].text = (" " * delta_spaces) + paragraph.runs[0].text
        else:
            greeting_run = paragraph.runs[greeting_run_idx]
            greeting_run.text = (" " * delta_spaces) + greeting_run.text
        return True

    remaining = abs(delta_spaces)
    for index in range(greeting_run_idx - 1, -1, -1):
        if remaining <= 0:
            break
        run = paragraph.runs[index]
        if run.text.strip():
            continue
        removable = min(remaining, len(run.text))
        run.text = run.text[: len(run.text) - removable]
        remaining -= removable

    return remaining == 0


def _create_variant(base_path: Path, target_path: Path, *, delta_spaces: int) -> None:
    shutil.copy2(base_path, target_path)
    doc = Document(str(target_path))
    if not _adjust_greeting_indent(doc, delta_spaces=delta_spaces):
        raise RuntimeError(f"Could not adjust greeting indent in {base_path.name}")
    doc.save(str(target_path))


def main() -> int:
    templates_dir = ROOT / "templates"
    base_path = templates_dir / TEMPLATE_DEFAULT
    if not base_path.exists():
        print(f"Base template not found: {base_path}")
        return 1

    long_path = templates_dir / TEMPLATE_LONG
    short_path = templates_dir / TEMPLATE_SHORT

    _create_variant(base_path, long_path, delta_spaces=-3)
    _create_variant(base_path, short_path, delta_spaces=3)

    print(f"Created: {long_path.name} (-3 spaces)")
    print(f"Created: {short_path.name} (+3 spaces)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
