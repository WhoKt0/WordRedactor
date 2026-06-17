"""Create sample Word template and Excel file for first run."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook


def create_letter_template(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_paragraph("Исх № {{OUT_NUMBER}} от {{DATE}}.")
    doc.add_paragraph("")
    doc.add_paragraph("Председателю правления")
    doc.add_paragraph("{{BANK_LEGAL_NAME}}")
    doc.add_paragraph("{{MR_MS}} {{CHAIR_SHORT_DATIVE}}")
    doc.add_paragraph("")
    doc.add_paragraph("{{GREETING_WORD}} {{GREETING_NAME}},")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Настоящим направляем коммерческое предложение RenSer Technologies "
        "по поставке банковских карт и PIN-конвертов."
    )
    doc.add_paragraph("")
    doc.add_paragraph("С уважением,")
    doc.add_paragraph("RenSer Technologies")

    doc.save(str(path))
    print(f"Created: {path}")


def create_sample_excel(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "banks"

    headers = [
        "bank_name",
        "bank_legal_name",
        "recipient_email",
        "cc_email",
        "bcc_email",
        "chair_full_name",
        "chair_short_dative",
        "gender",
        "greeting_name",
        "pdf_filename",
        "email_bank_name",
    ]
    ws.append(headers)

    ws.append(
        [
            "Алокабанк",
            'АК "Алокабанк"',
            "info@example.uz",
            "",
            "",
            "Ирисбекова Каммуна Наринбаевна",
            "Ирисбековой К. Н.",
            "female",
            "Каммуна Наринбаевна",
            "Предложение_по_картам_Алокабанк.pdf",
            'АК "Алокабанк"',
        ]
    )

    wb.save(path)
    print(f"Created: {path}")


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    create_letter_template(root / "templates" / "letter_template.docx")
    create_sample_excel(root / "data" / "banks.xlsx")


if __name__ == "__main__":
    main()
